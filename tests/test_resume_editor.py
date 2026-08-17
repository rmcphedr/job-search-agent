from pathlib import Path

from docx import Document

from src.documents.resume_editor import (
    docx_to_html,
    editable_resume_path,
    render_docx_preview,
    save_html_to_docx,
)


def test_resume_editor_round_trip_and_path_validation(tmp_path: Path) -> None:
    path = tmp_path / "resume.docx"
    document = Document()
    document.add_heading("Candidate Name", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run("Python").bold = True
    paragraph.add_run(" and research")
    document.add_paragraph("Built useful systems", style="List Bullet")
    document.save(path)

    html = docx_to_html(path)
    assert "<h1>Candidate Name</h1>" in html
    assert "<strong>Python</strong>" in html
    assert editable_resume_path(str(path)) == path.resolve()
    assert editable_resume_path(str(tmp_path / "missing.docx")) is None
    assert editable_resume_path(str(tmp_path / "resume.pdf")) is None

    save_html_to_docx(
        path,
        "<h1>Updated Name</h1><p><strong>AI</strong> and <em>healthcare</em></p>"
        "<ul><li>Shipped a product</li></ul>",
    )

    updated = Document(path)
    assert [paragraph.text for paragraph in updated.paragraphs] == [
        "Updated Name",
        "AI and healthcare",
        "Shipped a product",
    ]
    assert updated.paragraphs[1].runs[0].bold
    assert not updated.paragraphs[1].runs[1].italic
    assert updated.paragraphs[1].runs[2].italic


def test_save_rejects_non_docx(tmp_path: Path) -> None:
    path = tmp_path / "resume.txt"
    path.write_text("resume", encoding="utf-8")

    try:
        save_html_to_docx(path, "<p>Changed</p>")
    except ValueError as exc:
        assert "existing .docx" in str(exc)
    else:
        raise AssertionError("Expected ValueError")


def test_preview_reuses_cached_render(tmp_path: Path) -> None:
    path = tmp_path / "resume.docx"
    Document().save(path)
    import hashlib
    from tempfile import gettempdir

    fingerprint = hashlib.sha256(f"{path.resolve()}:{path.stat().st_mtime_ns}".encode()).hexdigest()[:16]
    preview = Path(gettempdir()) / "job-search-agent-previews" / fingerprint / "resume.pdf"
    preview.parent.mkdir(parents=True, exist_ok=True)
    preview.write_bytes(b"%PDF fixture")

    rendered, error = render_docx_preview(path)
    assert rendered == preview
    assert error is None
