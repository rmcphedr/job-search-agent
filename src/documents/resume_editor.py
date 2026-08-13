"""Convert resume DOCX bodies to and from editor-friendly HTML."""

from __future__ import annotations

from html import escape
from html.parser import HTMLParser
import hashlib
from pathlib import Path
import shutil
import subprocess
from tempfile import NamedTemporaryFile
from tempfile import gettempdir

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH


def editable_resume_path(value: str) -> Path | None:
    """Return a local DOCX path when the application field can be edited."""
    raw = str(value or "").strip()
    if not raw:
        return None
    path = Path(raw).expanduser()
    if path.suffix.lower() != ".docx" or not path.is_file():
        return None
    return path.resolve()


def docx_to_html(path: Path) -> str:
    """Render the document body as the small HTML subset supported by Quill."""
    document = Document(path)
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = "".join(_run_html(run) for run in paragraph.runs) or escape(paragraph.text)
        style = str(paragraph.style.name or "") if paragraph.style else ""
        if style.startswith("Heading"):
            level = style.removeprefix("Heading").strip() or "2"
            tag = f"h{level}" if level in {"1", "2", "3", "4", "5", "6"} else "h2"
        elif "List Bullet" in style:
            tag = "li"
            text = f"{text}"
        else:
            tag = "p"
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            text = f'<span class="ql-align-center">{text}</span>'
        blocks.append(f"<{tag}>{text or '<br>'}</{tag}>")
    return "".join(blocks)


def save_html_to_docx(path: Path, content: str) -> None:
    """Replace a DOCX body atomically while retaining its styles and sections."""
    path = path.expanduser().resolve()
    if path.suffix.lower() != ".docx" or not path.is_file():
        raise ValueError("Resume must be an existing .docx file.")

    document = Document(path)
    _clear_body(document)
    parser = _ResumeHTMLParser(document)
    parser.feed(content)
    parser.close()
    if not document.paragraphs:
        document.add_paragraph("")

    with NamedTemporaryFile(dir=path.parent, suffix=".docx", delete=False) as handle:
        temporary_path = Path(handle.name)
    try:
        document.save(temporary_path)
        temporary_path.replace(path)
    finally:
        temporary_path.unlink(missing_ok=True)


def render_docx_preview(path: Path) -> tuple[Path | None, str | None]:
    """Render a DOCX as a faithful PDF, or a Quick Look image on macOS."""
    path = path.expanduser().resolve()
    if not path.is_file():
        return None, "The document file no longer exists."
    fingerprint = hashlib.sha256(f"{path}:{path.stat().st_mtime_ns}".encode()).hexdigest()[:16]
    preview_dir = Path(gettempdir()) / "job-search-agent-previews" / fingerprint
    preview_dir.mkdir(parents=True, exist_ok=True)

    pdf_path = preview_dir / f"{path.stem}.pdf"
    if pdf_path.is_file():
        return pdf_path, None
    image_path = preview_dir / f"{path.name}.png"
    if image_path.is_file():
        return image_path, None

    libreoffice = shutil.which("soffice") or shutil.which("libreoffice")
    app_binary = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
    if not libreoffice and app_binary.is_file():
        libreoffice = str(app_binary)
    if libreoffice:
        completed = subprocess.run(
            [libreoffice, "--headless", "--convert-to", "pdf", "--outdir", str(preview_dir), str(path)],
            text=True,
            capture_output=True,
            timeout=60,
            check=False,
        )
        if completed.returncode == 0 and pdf_path.is_file():
            return pdf_path, None

    quicklook = shutil.which("qlmanage")
    if quicklook:
        completed = subprocess.run(
            [quicklook, "-t", "-s", "1600", "-o", str(preview_dir), str(path)],
            text=True,
            capture_output=True,
            timeout=60,
            check=False,
        )
        if completed.returncode == 0 and image_path.is_file():
            return image_path, None

    return None, "Install LibreOffice for an embedded multi-page PDF preview."


def _run_html(run) -> str:
    value = escape(run.text).replace("\n", "<br>")
    if run.bold:
        value = f"<strong>{value}</strong>"
    if run.italic:
        value = f"<em>{value}</em>"
    if run.underline:
        value = f"<u>{value}</u>"
    return value


def _clear_body(document: DocumentType) -> None:
    body = document._element.body
    for child in list(body):
        if not child.tag.endswith("sectPr"):
            body.remove(child)


class _ResumeHTMLParser(HTMLParser):
    BLOCKS = {"p", "div", "h1", "h2", "h3", "h4", "h5", "h6", "li"}

    def __init__(self, document: DocumentType) -> None:
        super().__init__(convert_charrefs=True)
        self.document = document
        self.paragraph = None
        self.bold = False
        self.italic = False
        self.underline = False

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in self.BLOCKS:
            style = None
            if tag.startswith("h") and tag[1:].isdigit():
                style = f"Heading {tag[1:]}"
            elif tag == "li":
                style = "List Bullet"
            try:
                self.paragraph = self.document.add_paragraph(style=style)
            except KeyError:
                self.paragraph = self.document.add_paragraph()
        elif tag in {"strong", "b"}:
            self.bold = True
        elif tag in {"em", "i"}:
            self.italic = True
        elif tag == "u":
            self.underline = True
        elif tag == "br":
            self._add_text("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in self.BLOCKS:
            self.paragraph = None
        elif tag in {"strong", "b"}:
            self.bold = False
        elif tag in {"em", "i"}:
            self.italic = False
        elif tag == "u":
            self.underline = False

    def handle_data(self, data: str) -> None:
        self._add_text(data)

    def _add_text(self, value: str) -> None:
        if not value:
            return
        if self.paragraph is None:
            if not value.strip():
                return
            self.paragraph = self.document.add_paragraph()
        run = self.paragraph.add_run(value)
        run.bold = self.bold
        run.italic = self.italic
        run.underline = self.underline
